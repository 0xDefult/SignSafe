import pdfplumber
from docx import Document
import io
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes
from fastapi import UploadFile, HTTPException

async def extract_text(file: UploadFile) -> str:
    """
    Extract raw text from PDF, DOCX, DOC, or Image files.
    Returns cleaned text string ready for Gemma analysis.
    Raises HTTPException if file type unsupported or extraction fails.
    """
    content = await file.read()
    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        return _extract_from_pdf(content)
    elif filename.endswith(".docx"):
        return _extract_from_docx(content)
    elif filename.endswith(".doc"):
        return _extract_from_doc(content)
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        return _extract_from_image(content)
    elif filename.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Please upload PDF, DOCX, DOC, or Image (PNG, JPG)."
        )

def _extract_from_pdf(content: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber with OCR fallback."""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- PAGE {page_num + 1} ---\n{page_text}")

        full_text = "\n\n".join(text_parts)

        # OCR Fallback for scanned PDFs
        if len(full_text.strip()) < 100:
            return _ocr_pdf(content)

        return full_text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF parsing failed: {str(e)}")

def _ocr_pdf(content: bytes) -> str:
    """Convert PDF to images and run OCR."""
    try:
        images = convert_from_bytes(content)
        text_parts = []
        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image)
            text_parts.append(f"--- PAGE {i+1} (OCR) ---\n{text}")
        return "\n\n".join(text_parts)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF OCR failed: {str(e)}")

def _extract_from_docx(content: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        doc = Document(io.BytesIO(content))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        return "\n\n".join(paragraphs)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX parsing failed: {str(e)}")

def _extract_from_doc(content: bytes) -> str:
    """
    Extract text from legacy .doc files.
    """
    try:
        # Basic attempt to decode, fallback to OCR
        try:
            return content.decode("utf-8", errors="ignore")
        except:
            return _extract_from_image(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOC parsing failed: {str(e)}")

def _extract_from_image(content: bytes) -> str:
    """Extract text from image bytes using Tesseract OCR."""
    try:
        image = Image.open(io.BytesIO(content))
        return pytesseract.image_to_string(image)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Image OCR failed: {str(e)}")
